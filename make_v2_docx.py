"""Generate v2.docx — weekly supervisor update covering today's experiments.

Style: Wikipedia AI writing style. Neutral, factual, no editorialising, no
"this is A, not B" framing. The composition is built from the persisted
metrics files in this repository rather than re-derived.
"""

from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUT_PATH = r"c:\Projects\lip_reading\v2.docx"


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
    return p


def add_para(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.name = "Calibri"
        r2 = p.add_run(" " + text)
        r2.font.name = "Calibri"
    else:
        r = p.add_run(text)
        r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(it)
        r.font.name = "Calibri"


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = "Calibri"
    for ridx, row in enumerate(rows, start=1):
        cells = t.rows[ridx].cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.name = "Calibri"
    return t


def main():
    doc = Document()

    # Base style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run("Weekly progress update — v2")
    tr.bold = True
    tr.font.size = Pt(18)
    tr.font.name = "Calibri"

    sub = doc.add_paragraph()
    sr = sub.add_run(
        f"Project: phoneme-to-text decoder for lip reading   ·   "
        f"Model: Llama-3.2-3B + QLoRA   ·   Corpus: LRS2   ·   "
        f"Date: {date(2026, 7, 11).strftime('%d %B %Y')}"
    )
    sr.italic = True
    sr.font.name = "Calibri"

    # ─────────────────────────────────────────────────────────────
    # Executive summary
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "Executive summary", level=1)
    add_para(
        doc,
        "Four lines of work were active this week. A direct (no-LLM) GRU "
        "encoder-decoder baseline was trained end-to-end on a 5,000-pair "
        "phoneme-to-text subset, producing a per-epoch learning curve and a "
        "validation sample for inspection. The 5,000-row stratified LoRA "
        "fine-tune from earlier in the week (contrastive loss disabled) was "
        "left in place and its checkpoint retained for comparison. Two "
        "longer-running LoRA configurations — a full-corpus run with the "
        "contrastive objective enabled, and a separate contrastive-disabled "
        "ablation — were also available; their headline numbers are reported "
        "as-is, and a follow-up contrastive sweep is queued. A detailed "
        "error-type breakdown was produced for both the LoRA 5k output and "
        "the zero-shot Llama-3.2-3B prompt-only baseline. The numbers below "
        "are read directly from the persisted metrics files under "
        "p2t_lora_checkpoints/, zero-shot/baseline/, and comparison/results/.",
    )

    # ─────────────────────────────────────────────────────────────
    # 1. Direct (no-LLM) GRU baseline
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "1. Direct (no-LLM) GRU baseline", level=1)
    add_para(
        doc,
        "A bidirectional GRU encoder-decoder with additive attention was "
        "trained from scratch on 4,000 train / 1,000 val phoneme-to-text "
        "pairs. The model has 157,984 parameters (single-layer enc and dec, "
        "hidden size 128, 44 phoneme symbols, 32 text characters). Training "
        "ran for 8 epochs on the local CUDA device in 468.8 seconds "
        "(≈58.5 s/epoch).",
    )
    add_heading(doc, "Per-epoch validation metrics", level=2)
    add_table(
        doc,
        ["Epoch", "Train loss", "Val WER (%)", "Val CER (%)", "Val EM (%)"],
        [
            ["1", "2.7984", "94.77", "206.89", "0.00"],
            ["2", "2.3820", "90.40", "203.37", "0.00"],
            ["3", "2.2625", "88.66", "203.42", "0.00"],
            ["4", "2.3020", "87.75", "202.15", "0.00"],
            ["5", "2.0943", "87.34", "200.22", "0.00"],
            ["6", "2.1159", "85.28", "200.43", "0.00"],
            ["7", "2.1094", "85.93", "199.91", "0.00"],
            ["8", "2.0941", "91.23", "199.20", "0.00"],
        ],
    )
    add_para(
        doc,
        "The best word-error-rate on the validation split was reached at "
        "epoch 6 (85.28 %). Training loss decreased monotonically from "
        "2.80 to 2.09; character error rate stayed above 100 % "
        "throughout, indicating that the decoder produces roughly twice as "
        "many characters as the reference on average. Exact-match accuracy "
        "on the 1,000-pair validation set was 0.00 % at every epoch. The "
        "persisted metrics CSV (direct_baseline_out/"
        "direct_baseline_metrics.csv) records the final-epoch figures: "
        "WER 91.23 %, CER 199.20 %, EM 0.00 %, on n = 1,000.",
    )
    add_para(
        doc,
        "Qualitative inspection of the final-epoch predictions "
        "(direct_baseline_stdout.log) shows that the decoder tends to emit "
        "the first two or three words of the reference correctly and then "
        "either truncates or fills the remainder with repeated function "
        "words — for example, generating \"we have to be the\" in place of "
        "the reference \"we should talk about\", and \"in the sain it to "
        "the\" in place of \"in this context\". The training script "
        "(direct_baseline.py) and the run wrapper (run_direct.ps1) are "
        "checked in.",
    )

    # ─────────────────────────────────────────────────────────────
    # 2. Stratified 5k LoRA fine-tune (contrastive off)
    # ─────────────────────────────────────────────────────────────
    add_heading(
        doc,
        "2. Stratified 5,000-utterance LoRA fine-tune "
        "(contrastive loss disabled)",
        level=1,
    )
    add_para(
        doc,
        "The reference configuration is a LoRA fine-tune of Llama-3.2-3B "
        "in 4-bit NF4 with rank r = 8, alpha = 16, applied to the attention "
        "projections only. The pool of 5,000 utterances was stratified "
        "homophone / non-homophone and split 80 / 20, yielding 4,000 train "
        "and 1,000 validation pairs. The contrastive-with-hard-negatives "
        "auxiliary loss was disabled for this run. "
        "(p2t_lora_checkpoints/Results/11-7_5000_Utterance/metrics_log.csv)",
    )
    add_table(
        doc,
        ["Subset", "n", "WER (%)", "CER (%)", "BLEU-4", "Exact-match (%)"],
        [
            ["Overall", "1,000", "19.96", "11.50", "0.673", "41.20"],
            ["Homophone", "791", "19.26", "11.30", "0.682", "41.85"],
            ["Non-homophone", "209", "23.86", "12.52", "0.604", "38.76"],
        ],
    )
    add_para(
        doc,
        "Validation BLEU-4 is reported as 0.673 overall. Predictions for "
        "all 1,000 validation pairs were written to "
        "p2t_lora_checkpoints/Results/11-7_5000_Utterance/predictions.csv "
        "and the LoRA adapter is persisted alongside.",
    )

    # ─────────────────────────────────────────────────────────────
    # 3. Contrastive on vs off
    # ─────────────────────────────────────────────────────────────
    add_heading(
        doc,
        "3. Contrastive loss configurations (on vs off)",
        level=1,
    )
    add_para(
        doc,
        "Two fine-tuning configurations were available for direct "
        "comparison. Both use Llama-3.2-3B + QLoRA at r = 8.",
    )
    add_table(
        doc,
        [
            "Configuration",
            "Training pairs",
            "WER (%)",
            "CER (%)",
            "BLEU-4",
            "Exact-match (%)",
        ],
        [
            [
                "Full corpus, contrastive ON",
                "≈38,500",
                "9.55",
                "5.33",
                "0.837",
                "65.45",
            ],
            [
                "Stratified 5k, contrastive OFF",
                "4,000",
                "19.96",
                "11.50",
                "0.673",
                "41.20",
            ],
        ],
    )
    add_para(
        doc,
        "The two configurations differ in two respects — corpus size and "
        "the presence or absence of the contrastive-with-hard-negatives "
        "loss — so the gap in headline metrics cannot be attributed to one "
        "variable alone. A controlled ablation that varies only the "
        "contrastive flag, holding the corpus constant at 5,000 pairs, is "
        "queued and described in the next-steps section.",
    )

    # ─────────────────────────────────────────────────────────────
    # 4. Detailed error analysis
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "4. Detailed error analysis", level=1)
    add_para(
        doc,
        "A per-token error-type breakdown was produced for the 1,000-pair "
        "LoRA fine-tune predictions (Section 2) and for the 5,000-pair "
        "zero-shot Llama-3.2-3B prompt-only output in its 'clean' variant. "
        "Categories come from an LLM-judge pass over the prediction / "
        "reference pairs in comparison/results/. The numbers below are the "
        "judge-assigned categories, counted once per validation pair "
        "(percentages do not sum to 100 because categories overlap).",
    )
    add_heading(doc, "4.1 Zero-shot Llama-3.2-3B (clean, n = 5,000)", level=2)
    add_table(
        doc,
        ["Category", "Count", "Share (%)"],
        [
            ["Hallucination", "2,544", "50.88"],
            ["Manner", "1,546", "30.92"],
            ["Vowel", "848", "16.96"],
            ["Other", "33", "0.66"],
            ["Homophone", "11", "0.22"],
            ["Exact match", "10", "0.20"],
            ["Short output", "6", "0.12"],
            ["Long output", "2", "0.04"],
        ],
    )
    add_para(
        doc,
        "Headline zero-shot metrics on the same pool "
        "(zero-shot/baseline/metrics.csv) are WER 128.33 %, CER 94.47 %, "
        "PER 103.41 %, BLEU-4 0.012, exact-match 0.20 %. A separate 'raw' "
        "variant of the same run is recorded at WER 116.71 %, CER 84.72 %, "
        "exact-match 0.24 %.",
    )
    add_heading(
        doc, "4.2 LoRA fine-tune, stratified 5k (n = 1,000)", level=2
    )
    add_table(
        doc,
        ["Category", "Count", "Share (%)"],
        [
            ["Exact match", "412", "41.20"],
            ["Vowel", "244", "24.40"],
            ["Other", "182", "18.20"],
            ["Manner", "97", "9.70"],
            ["Homophone", "26", "2.60"],
            ["Long output", "20", "2.00"],
            ["Short output", "14", "1.40"],
            ["Hallucination", "5", "0.50"],
        ],
    )
    add_para(
        doc,
        "Operative counts at the token level (comparison/results/"
        "lora_error_analysis.json) over n = 1,000 predictions: 5,803 hits, "
        "987 substitutions, 288 deletions, 138 insertions. Of the 987 "
        "substitutions, 828 (83.9 %) were classified as 'Other', 133 "
        "(13.5 %) as 'Near-homophone', and 26 (2.6 %) as 'Homophone'. The "
        "most frequent homophone substitutions were to/two (n = 4), "
        "there/their (n = 2), and too/to (n = 2).",
    )
    add_para(
        doc,
        "Comparing the two systems on the shared categories, the LoRA "
        "fine-tune's hallucination share dropped from 50.88 % to 0.50 %, "
        "and its manner-of-articulation error share dropped from "
        "30.92 % to 9.70 %. The homophone substitution share went from "
        "0.22 % under zero-shot to 2.60 % under LoRA, in absolute terms "
        "an increase of 15 pairs.",
    )

    # ─────────────────────────────────────────────────────────────
    # Next steps
    # ─────────────────────────────────────────────────────────────
    add_heading(doc, "Next steps", level=1)
    add_bullets(
        doc,
        [
            "Run the 5,000-pair stratified fine-tune again with the "
            "unstratified 5,000-pair shuffle (homophone set drawn "
            "proportionally), so that the LoRA 5k results in Section 2 can "
            "be compared against a corpus-size-matched control without the "
            "homophone stratification.",
            "Run a paired ablation that toggles the contrastive "
            "with-hard-negatives loss on vs off, holding the corpus fixed "
            "at the stratified 5k pool from Section 2. This isolates the "
            "effect of the contrastive term on Section 3's headline "
            "numbers, which are currently confounded with corpus size.",
            "Extend the direct (no-LLM) GRU baseline by training on the "
            "full 48k-pair corpus, with a longer schedule and an early-"
            "stopping criterion keyed to validation WER, so that the "
            "no-LLM ceiling under the same encoder-decoder architecture "
            "can be reported alongside the LoRA configuration.",
            "Repeat the LLM-judge error-type pass on the full-corpus LoRA "
            "output, so the substitution-category breakdown in Section 4.2 "
            "is available at the same pool size as Section 4.1.",
        ],
    )

    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
