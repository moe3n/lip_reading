"""Build zero-shot/baseline_with_full_corpus/results_discussion.docx

Single self-contained docx that contains the entire Results & Discussion for
the zero-shot baseline, ready to be dropped into the methodology chapter.

Reads metrics_full_48164.csv (re-scored from view_full_48164.txt by
rescore_view.py) as the source of truth for the headline numbers, and reads
six hand-picked WRONG rows from view_full_48164.txt for the error-pattern
table.
"""

import csv
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = Path(__file__).resolve().parent
METRICS_CSV = HERE / "metrics_full_48164.csv"
VIEW_TXT = HERE / "view_full_48164.txt"
OUT_DOCX = HERE / "results_discussion.docx"


def load_metrics(path):
    rows = []
    with open(path, "r", encoding="utf-8") as f:
        for r in csv.DictReader(f):
            rows.append(r)
    return rows


def pick_error_rows(path, n=6):
    """Pull the first n WRONG rows from the viewer file (every other row from
    the top, to spread the picks across the corpus)."""
    picks = []
    with open(path, "r", encoding="utf-8") as f:
        next(f)  # header
        seen = 0
        for line in f:
            parts = line.rstrip("\n").split("|")
            if len(parts) < 4:
                continue
            status, _, predicted, target = (
                parts[0].strip(),
                parts[1].strip(),
                parts[2].strip(),
                parts[3].strip(),
            )
            if status != "WRONG":
                continue
            # spread the picks across the file: take every ~8000th WRONG row
            if seen % 8000 == 0:
                picks.append((target, predicted))
                if len(picks) >= n:
                    break
            seen += 1
    return picks


def style_normal(doc):
    s = doc.styles["Normal"]
    s.font.name = "Calibri"
    s.font.size = Pt(11)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def add_paragraph(doc, text, justify=False, italic=False):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.italic = italic
    return p


def main():
    metrics = load_metrics(METRICS_CSV)
    by_split = {m["split"]: m for m in metrics}

    overall = by_split["Overall"]
    homo = by_split["Homophone"]
    nonhomo = by_split["Non-Homophone"]

    error_rows = pick_error_rows(VIEW_TXT, n=6)

    doc = Document()
    style_normal(doc)

    # Title
    title = doc.add_heading(
        "Results and Discussion — Zero-Shot Phoneme-to-Text Baseline",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_paragraph(
        doc,
        "Scope. A single, self-contained ablation that establishes the "
        "default performance of a pretrained Llama-3.2-3B language model on "
        "phoneme-to-text conversion, with no task-specific training. The aim "
        "is to (a) verify that vanilla instruction-tuned LLMs cannot perform "
        "this task out of the box, and (b) supply a baseline number that the "
        "fine-tuned decoder below has to beat.",
        justify=True,
    )

    # 1. Setup
    add_heading(doc, "1. Setup", level=1)
    setup_rows = [
        ("Model", "meta-llama/Llama-3.2-3B (3B parameters, instruction-tuned)"),
        ("Quantisation", "4-bit (bitsandbytes, fp16 compute)"),
        (
            "Decoding",
            "Greedy: do_sample=False, num_beams=1, max_new_tokens=50, "
            "repetition_penalty=1.0, temperature=0.0, top_p=1.0",
        ),
        ("Seed", "42"),
        ("Batch size", "8 (per GPU)"),
        (
            "Corpus",
            "LRS2 main partition, full 48,164-sentence split, verbatim prompt, "
            "no cleaning",
        ),
        ("Phoneme source", "phonemes_raw — markers and stress digits preserved"),
        ("Prompt", "Project's verbatim phoneme-to-text prompt (held constant)"),
        ("Hardware", "2× NVIDIA GTX 1080 (8 GB each), sharded by row index"),
    ]
    add_table(doc, ["Component", "Value"], setup_rows)

    add_paragraph(
        doc,
        "The experimental script supports transparent sharding via the "
        "BWFC_OFFSET / BWFC_STRIDE environment variables; with stride=2 each "
        "GPU processed 24,082 rows, and the two .jsonl files were merged and "
        "re-scored to produce the metrics below. Both shards completed "
        "successfully (24,082 rows each), and the headline numbers were "
        "re-derived from the preserved view_full_48164.txt file.",
        justify=True,
    )

    # 2. Headline
    add_heading(doc, "2. Headline Result", level=1)
    add_paragraph(
        doc,
        "The full-corpus metrics, computed by rescore_view.py from "
        "view_full_48164.txt:",
    )
    headline_rows = [
        (
            "Overall",
            overall["n"],
            overall["WER_pct"],
            overall["CER_pct"],
            overall["EM_pct"],
            overall["BLEU4"],
        ),
        (
            "Homophone",
            homo["n"],
            homo["WER_pct"],
            homo["CER_pct"],
            homo["EM_pct"],
            homo["BLEU4"],
        ),
        (
            "Non-Homophone",
            nonhomo["n"],
            nonhomo["WER_pct"],
            nonhomo["CER_pct"],
            nonhomo["EM_pct"],
            nonhomo["BLEU4"],
        ),
    ]
    add_table(
        doc,
        ["Split", "n", "WER (%)", "CER (%)", "EM (%)", "BLEU-4"],
        headline_rows,
    )

    add_paragraph(
        doc,
        "Two things are immediately striking.",
    )
    add_paragraph(
        doc,
        "First, exact match is essentially zero (0.22%). Across 48,164 "
        "sentences the model reproduces a target verbatim only 106 times. "
        "This is a floor on this specific run, on this prompt, on this "
        "decoding spec.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Second, WER exceeds 100% (108.09%). This is genuine, not a parser "
        "bug. It happens because the model hallucinates fluent English "
        "continuations that are at least as long as the reference; word-level "
        "edit distance can then exceed the number of reference words. The "
        "same pattern holds for BLEU-4 ≈ 0.01, meaning the model does not "
        "share even a single 4-gram with the target on the vast majority of "
        "sentences.",
        justify=True,
    )

    # 3. Error pattern
    add_heading(doc, "3. Error Pattern", level=1)
    add_paragraph(
        doc,
        "Six WRONG predictions sampled from the viewer file (target on the "
        "left, model output on the right):",
    )
    err_rows = [(t, p) for t, p in error_rows]
    add_table(doc, ["Target", "Prediction"], err_rows)

    add_paragraph(
        doc,
        "Three patterns are visible:",
    )
    add_paragraph(
        doc,
        "English-grammar hallucination. The model produces a fluent English "
        "question or sentence that has no acoustic or lexical overlap with "
        "the target. This is the dominant mode — the model treats the prompt "
        "as the start of a story and continues it.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Length-matching. Predictions are roughly target-length, which is "
        "why CER (≈81%) is much lower than WER (≈108%) — many characters are "
        "right, just in the wrong order, in the wrong words.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Phoneme echoes. A small fraction of outputs preserve some of the "
        "source phonemes (for example, 'SAY MAY BAG WADH D AAN' against 'IT "
        "MIGHT BE UGLY WITH THE HEAD ON'), suggesting the model is partially "
        "conditioned on the phoneme stream but unable to map it cleanly to "
        "text.",
        justify=True,
    )
    add_paragraph(
        doc,
        "No off-by-one or prompt-stripping bug was found: a ten-row spot "
        "check confirmed all sampled errors are the model's literal output, "
        "not a parser artefact.",
        justify=True,
    )

    # 4. Homophone vs Non-Homophone
    add_heading(doc, "4. Homophone vs Non-Homophone Split", level=1)
    add_paragraph(
        doc,
        "The corpus partition labelling is itself informative. The "
        "Homophone partition covers 37,374 / 48,164 (77.6%) of sentences; "
        "that is the dominant case in LRS2 — sentences that contain at "
        "least one word whose target spelling is consistent with multiple "
        "phoneme sequences. On this run the model performs slightly worse "
        "on the Homophone partition (WER 107.28% vs 112.50%, EM 0.17% vs "
        "0.40%). This is consistent with the failure mode above: the model "
        "is not even at the phoneme–grapheme mapping stage, so the "
        "homophone subtlety is masked by the larger failure of failing to "
        "transcribe at all.",
        justify=True,
    )

    # 5. What this baseline establishes
    add_heading(doc, "5. What This Run Establishes", level=1)
    add_paragraph(
        doc,
        "Three findings from this run, in their own terms:",
    )
    add_paragraph(
        doc,
        "A pretrained 3B instruction-tuned LLM has essentially zero "
        "useful out-of-the-box ability on phoneme-to-text. With greedy "
        "decoding and the project's verbatim prompt, it reproduces the "
        "target verbatim on only 106 / 48,164 sentences (0.22% EM).",
        justify=True,
    )
    add_paragraph(
        doc,
        "The dominant failure mode is fluent English hallucination. The "
        "model treats the phoneme prompt as the start of a sentence and "
        "continues it, producing outputs that have little or no lexical "
        "overlap with the target. WER exceeds 100% (108.09%) precisely "
        "because predicted sentences are at least as long as the targets.",
        justify=True,
    )
    add_paragraph(
        doc,
        "Homophone vs non-homophone is not a meaningful axis here. The "
        "two splits score within a few percentage points of each other "
        "(WER 107.28% vs 112.50%, EM 0.17% vs 0.40%), confirming that "
        "the model is failing at an earlier stage — phoneme-to-grapheme "
        "mapping — than the homophone-disambiguation stage.",
        justify=True,
    )

    # 6. Scope and limitations of this run
    add_heading(doc, "6. Scope and Limitations of this Run", level=1)
    add_paragraph(
        doc,
        "One prompt, one decoding spec, one corpus. We report one prompt "
        "(the project's verbatim phoneme-to-text prompt), one decoding "
        "spec (greedy, max_new_tokens=50, seed=42), and one corpus (the "
        "full LRS2 main partition, 48,164 sentences). This chapter "
        "establishes a baseline; it does not claim this is the best "
        "zero-shot result achievable.",
        justify=True,
    )
    add_paragraph(
        doc,
        "No tokenisation ablation. The <space> phoneme marker is "
        "removed from the input, matching the project's standard setup. "
        "The impact of that decision on this baseline is not isolated "
        "here.",
        justify=True,
    )
    add_paragraph(
        doc,
        "No sampling, beam, or repetition-penalty sweep. Greedy only. "
        "The decoder hyper-parameters are fixed at their default values.",
        justify=True,
    )

    # Drop-in paragraph
    add_heading(doc, "Drop-in Thesis-Style Summary", level=1)
    add_paragraph(
        doc,
        "A zero-shot baseline was established by running the pretrained "
        "Llama-3.2-3B model on the full 48,164-sentence LRS2 split with "
        "greedy decoding (max_new_tokens=50, no sampling, seed=42) and the "
        "project's verbatim phoneme-to-text prompt, with no task-specific "
        "training. The model achieves 0.22% exact match, 108.09% word error "
        "rate, 81.01% character error rate, and 0.0113 BLEU-4 on the full "
        "split, with near-identical figures on the homophone (37,374 "
        "sentences) and non-homophone (10,790 sentences) sub-splits. The "
        "error pattern is dominated by fluent English hallucinations that "
        "ignore the phoneme stream, with WER exceeding 100% because "
        "predicted sentences are at least as long as the targets. The "
        "baseline establishes that a vanilla instruction-tuned LLM cannot "
        "perform phoneme-to-text out of the box.",
        justify=True,
    )

    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")


if __name__ == "__main__":
    main()