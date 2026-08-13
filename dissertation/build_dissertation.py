"""Build the dissertation .docx to the LSBU MSc dissertation guide's format:
Times New Roman 12pt, double spacing, 1.5in binding (left) margin and 0.75in
elsewhere, title-page fields per section 5.1.3.1, and page numbers in the top-right
corner numbered consecutively (title page suppressed). Each chapter starts on a new
page. Body content for Chapters 2 and 3 comes from chapters_2_3.md.
"""

import os
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "analysis"))
from md_to_docx import render_markdown_into

HERE = os.path.dirname(os.path.abspath(__file__))
BODY_MD = [os.path.join(HERE, f) for f in (
    "chapter_1.md", "chapters_2_3.md", "chapter_4.md", "chapter_5.md",
    "chapter_6.md", "chapter_7.md", "chapter_8.md",
    "chapter_9.md", "chapter_10.md", "chapter_11.md", "chapter_12.md", "references.md")]
OUT = os.path.join(HERE, os.environ.get("DISS_OUT", "dissertation_draft.docx"))

# Placeholders the student fills in. Kept obvious so none is missed.
TITLE = "Phoneme-to-Text Conversion for Automated Lip-Reading using a Decoder-Only Large Language Model with QLoRA Fine-Tuning and Noise-Augmented Robustness"
AUTHOR = "[Surname, Initial]"
STUDENT_ID = "[Student Number]"
SUPERVISOR = "Mr Daqing Chen"
DEGREE = "MSc [Programme Name]"
# Affiliation: the two LSBU guidelines disagree; confirm which the current template uses.
DIVISION = "Division of Computer Science and Informatics, School of Engineering"
DIVISION_ALT = "(or: School of Computer Science and Digital Technologies -- confirm current name)"
UNIVERSITY = "London South Bank University"
DATE = "[Month Year]"


def add_top_right_page_number(section):
    """Consecutive page number in the top-right corner (header), per 5.1.2.4."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    b = OxmlElement("w:fldChar"); b.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    e = OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"), "end")
    run._r.append(b); run._r.append(instr); run._r.append(e)
    run.font.name = "Times New Roman"; run.font.size = Pt(12)


def centered(doc, text, size, bold=False, space_after=12, space_before=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE  # title page single-spaced
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def main():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after = Pt(0)

    # Headings also Times New Roman.
    for hs in ("Heading 1", "Heading 2", "Heading 3"):
        try:
            doc.styles[hs].font.name = "Times New Roman"
        except KeyError:
            pass

    sec = doc.sections[0]
    sec.left_margin = Inches(1.5)      # binding edge
    sec.right_margin = Inches(0.75)
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    # Consecutive Arabic numbers, top-right, title page suppressed.
    sec.different_first_page_header_footer = True
    add_top_right_page_number(sec)

    # ── Title page (page 1, number suppressed) ────────────────────────────────
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    centered(doc, UNIVERSITY, 16, bold=True, space_after=4)
    centered(doc, DIVISION, 12, space_after=2)
    centered(doc, DIVISION_ALT, 10, space_after=40)
    centered(doc, TITLE, 16, bold=True, space_after=40)
    centered(doc, "A dissertation submitted in partial fulfilment of the requirements for the degree of", 12, space_after=4)
    centered(doc, DEGREE, 12, bold=True, space_after=40)
    centered(doc, f"{AUTHOR}", 12, space_after=4)
    centered(doc, f"Student Number: {STUDENT_ID}", 12, space_after=4)
    centered(doc, f"Supervisor: {SUPERVISOR}", 12, space_after=24)
    centered(doc, DATE, 12)
    page_break(doc)

    # ── Preliminaries (each on a new page) ────────────────────────────────────
    # Abstract is rendered from abstract.md; the rest stay as placeholders.
    doc.add_heading("Abstract", level=1)
    abstract_path = os.path.join(HERE, "abstract.md")
    if os.path.isfile(abstract_path):
        with open(abstract_path, encoding="utf-8") as f:
            doc.add_paragraph(" ".join(l.strip() for l in f if l.strip()))
    else:
        doc.add_paragraph("[To be completed.]")
    page_break(doc)
    for title in ["Acknowledgements", "Table of Contents", "List of Figures",
                  "List of Tables", "List of Abbreviations"]:
        doc.add_heading(title, level=1)
        doc.add_paragraph("[To be completed.]")
        page_break(doc)

    # ── Body ──────────────────────────────────────────────────────────────────
    for md in BODY_MD:
        render_markdown_into(doc, md)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
