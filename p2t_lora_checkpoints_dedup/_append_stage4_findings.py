"""Append Sections 11-13 to findings.docx documenting the Stage 4
grammar / semantic / consolidated-summary results.

Stage 4 inputs (already produced by step4_grammar_semantic.py and
stage4_summary.py):

    analysis/tables/grammar_breakdown.csv      (per-EM-False-row resolution)
    analysis/tables/casing_punct_audit.csv     (casing/punct/digit buckets)
    analysis/tables/semantic_similarity.csv    (per-EM-False-row BERTScore F1)
    analysis/tables/semantic_refs.csv          (refs side for reruns)
    analysis/tables/semantic_hyps.csv          (hyps side for reruns)
    analysis/tables/stage4_metrics_summary.json
    analysis/tables/priority_ranked.csv
    analysis/tables/priority_summary.json

Stage 4 figures (already produced by _make_stage4_figures.py):
    analysis/figures/fig16_bertscore_distribution.png
    analysis/figures/fig17_substitution_categories.png
    analysis/figures/fig18_priority_ranked.png

Run once. Will duplicate if re-run on a doc that already has the
appended sections; safe to delete the appended block first.
"""

import csv
import json
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
TABLES = ROOT / "analysis" / "tables"
FIGURES = ROOT / "analysis" / "figures"
DOC = ROOT / "analysis" / "findings.docx"


# --- helpers (mirror _make_findings_docx.py) -------------------------------

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_figure(doc, png_path, caption, width_inches=6.5):
    if not Path(png_path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(png_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    return t


def read_csv(p):
    with p.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def read_json(p):
    with p.open(encoding="utf-8") as f:
        return json.load(f)


def fmt_pct(x):
    return f"{float(x):.2f}%"


# --- build -----------------------------------------------------------------

doc = Document(DOC)
gram_summary = read_json(TABLES / "stage4_metrics_summary.json")
priority = read_csv(TABLES / "priority_ranked.csv")

# ---------- Section 11: Grammar (Step 7, Option 3) --------------------------

add_heading(doc, "11. Grammar analysis (closed-class dep-role mismatch)", level=1)
add_para(doc,
    "Step 7 of the extended methodology. The canonical Mira Fleite tool "
    "is `language_tool_python`, which requires a local JRE; no JRE is "
    "installed in this environment. We use the built-but-unapplied "
    "fallback: `contextual_analysis.check_grammar` (Option 3, pure-Python "
    "spaCy). check_grammar indexes into the hypothesis sentence, runs "
    "spaCy dependency parsing, and flags substitutions where the "
    "predicted word is one of {THEIR, YOUR, ITS, MY, OUR, WHOSE} but "
    "isn't syntactically restricted to a `poss` (possessive) role while "
    "the reference required it."
)
add_para(doc,
    "We invoke Option 3 via `error_analysis.analyze_pair(use_llm=False)` "
    "which routes every substitution through `classify_substitution` "
    "and then escalates Homophone (and, before the speed-patch, "
    "Near-homophone) substitutions through check_grammar. The "
    "near-homophone path is a brute-force scan over the ~125k-word CMU "
    "dictionary and would have added several minutes of runtime at full "
    "corpus scale; we patched it to an O(1)-only homophone lookup "
    "(preserving the closed-class detector's exact-match coverage of "
    "{THEIR/THERE, YOUR/YORE, ITS/IT'S, ...}). The cost is that we no "
    "longer classify Near-homophone pairs as candidates for Option 3 "
    "escalation, which is acceptable here because the closed-class set "
    "covers ~all role-mismatch confusions of interest."
)
add_para(doc,
    "We also run a separate cheap mechanical audit on all 949 hypotheses "
    "to record casing, trailing-period and digit-token issues, since the "
    "corpus has eyeballed casing artefacts (uppercase words in "
    "otherwise-uniform hyp text)."
)

gram_rows = read_csv(TABLES / "grammar_breakdown.csv")
em_f_rows = [r for r in gram_rows if r["exact_match"] == "False"]
total_subs = sum(int(r["n_subs"]) for r in gram_rows)
em_f_subs = sum(int(r["n_subs"]) for r in em_f_rows)
homophone_subs = sum(int(r["homophone_subs"] or 0) for r in em_f_rows)
other_subs = sum(int(r["other_subs"] or 0) for r in em_f_rows)
resolved_by_grammar = gram_summary["step7_grammar"]["n_resolved_by_grammar"]

add_table(
    doc,
    ["Group", "n", "# substitutions", "# Homophone subs", "# Other subs"],
    [
        ["All 949 rows", len(gram_rows), str(total_subs), "-", "-"],
        ["EM-False (76)", len(em_f_rows), str(em_f_subs), str(homophone_subs), str(other_subs)],
    ],
)

audit = read_csv(TABLES / "casing_punct_audit.csv")
add_table(
    doc,
    ["Flag", "EM-True", "EM-False"],
    [[r["flag"], r["n"], "-"] if r["exact_match"] == "True"
        else [r["flag"], "-", r["n"]]
     for r in audit],
)

add_para(doc, "Reading:", bold=True)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "The 4 Homophone substitutions on the EM-False slice are "
    "(TO/TOO), (AD/ADD), (BY/BUY), (LLOYD/LOYD). None of them involves a "
    "closed-class possessive pronoun, so the Option 3 grammar detector "
    "fires 0 times on this corpus."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Of the 85 substitutions scored on the EM-False slice, 81 (95.3%) "
    "fall into the 'Other' bucket — meaning they aren't covered by the "
    "exact-homophone CMU lookup. These are the multi-word, hyphenation, "
    "suffixation, and OOV-type errors that we already saw in the manual "
    "audit (Section 4)."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "The mechanical audit sees 932/949 hyps with at least one uppercase "
    "word — this is by design (the corpus is uppercase-by-convention); "
    "17/949 contain a digit. None of these is a real decoding failure "
    "but they do inflate CER; we already report CER (0.98%) separately "
    "from WER (2.09%) in the headline."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Take-away: the grammar detector is correctly armed and would fire "
    "if the corpus contained (THEIR/THERE) or (YOUR/YORE) "
    "role-misassignment, but the dedup beam-5 corpus does not contain "
    "any such pairs. Keep check_grammar armed for production-scale runs."
)

add_figure(doc, FIGURES / "fig17_substitution_categories.png",
    "Figure 17 — Substitution categories on the EM-False slice. 'Other' "
    "dominates; the closed-class grammar path contributes nothing on this "
    "corpus, consistent with the absence of possessive-vs-location "
    "homophone confusions in dedup val.")


# ---------- Section 12: Semantic similarity (Step 8, BERTScore) ------------

add_heading(doc, "12. Semantic similarity (BERTScore F1, EM-False slice)", level=1)
add_para(doc,
    "Step 8 of the extended methodology. BERTScore F1 between reference "
    "and hypothesis on the 76 EM-False rows, computed with "
    "`microsoft/deberta-base-mnli` (smaller and faster than the default "
    "roberta-large; sufficient for English on a CPU box). Layer 10 "
    "(default for deberta-base)."
)

sem_rows = read_csv(TABLES / "semantic_similarity.csv")
f1_list = [float(r["bertscore_f1"]) for r in sem_rows]
mean_f1 = sum(f1_list) / len(f1_list)
median_f1 = sorted(f1_list)[len(f1_list) // 2]
above_090 = sum(1 for x in f1_list if x >= 0.90)
above_070 = sum(1 for x in f1_list if x >= 0.70)
below_050 = sum(1 for x in f1_list if x < 0.50)

homo_f1 = [float(r["bertscore_f1"]) for r in sem_rows
            if r["is_homophone"] == "True"]
nonhomo_f1 = [float(r["bertscore_f1"]) for r in sem_rows
               if r["is_homophone"] == "False"]

add_table(
    doc,
    ["Statistic", "Value"],
    [
        ["n EM-False rows", str(len(f1_list))],
        ["model", "microsoft/deberta-base-mnli, layer 10"],
        ["mean F1", f"{mean_f1:.4f}"],
        ["median F1", f"{median_f1:.4f}"],
        [f"F1 \u2265 0.90 (semantically-preserved)", f"{above_090}/{len(f1_list)} ({100.0 * above_090 / len(f1_list):.1f}%)"],
        [f"F1 \u2265 0.70 (broad-match)", f"{above_070}/{len(f1_list)} ({100.0 * above_070 / len(f1_list):.1f}%)"],
        ["F1 < 0.50 (semantic mismatch)",
            f"{below_050}/{len(f1_list)} ({100.0 * below_050 / len(f1_list):.1f}%)"],
        ["homophone rows mean F1",
            f"{sum(homo_f1)/len(homo_f1):.4f}  (n={len(homo_f1)})" if homo_f1 else "n/a"],
        ["non-homophone rows mean F1",
            f"{sum(nonhomo_f1)/len(nonhomo_f1):.4f}  (n={len(nonhomo_f1)})" if nonhomo_f1 else "n/a"],
    ],
)

add_para(doc, "Reading:", bold=True)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    f"On the EM-False slice (n=76), the model's mean BERTScore F1 is "
    f"{mean_f1:.4f} — even when WER is non-zero, the hypothesis is "
    "semantically close to the reference. "
    f"{above_090}/{len(f1_list)} ({100.0*above_090/len(f1_list):.1f}%) "
    "of EM-False rows score \u2265 0.90."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Homophone and non-homophone rows score within ~1 F1-point of each "
    "other (homophone slightly higher). This is consistent with the "
    "phoneme-level AER breakdown in Section 8: most errors are "
    "phonetically close to the reference, so the BERT contextual "
    "embeddings remain near-identical."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    f"Only {below_050} row(s) fall below 0.50. These are the genuinely "
    "different outputs — e.g. the manual audit row 'CHILD OF THE "
    "MILLENNIUM' \u2192 'CHILD OF THE MOLYNEUX', where the model "
    "substituted an unrelated surname."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Reporting implication: WER (and especially CER) for lip-reading "
    "overstates the practical error magnitude. Reporting WER/CER "
    "alongside BERTScore and a manual audit gives a more honest picture "
    "than any one number."
)

add_figure(doc, FIGURES / "fig16_bertscore_distribution.png",
    "Figure 16 — BERTScore F1 distribution on the 76 EM-False rows. "
    "Vertical dashed lines mark the 0.90 and 0.70 thresholds; dotted "
    "line is the mean (0.8839).")


# ---------- Section 13: Consolidated priority-ranked summary ---------------

add_heading(doc, "13. Consolidated priority-ranked summary", level=1)
add_para(doc,
    "Step 10 of the extended methodology. Reads every Stage 1-4 table, "
    "ranks the issues by impact and remediation cost, and emits "
    "`priority_ranked.csv` and `priority_summary.json`. The ranking is "
    "principle-based (not score-based), reflecting whether the issue "
    "touches many rows, has a clear scope, and admits a mechanical fix."
)

add_table(
    doc,
    ["Rank", "Issue", "Evidence", "Recommended fix (one-liner)"],
    [[r["rank"], r["issue"][:60] + ("..." if len(r["issue"]) > 60 else ""),
      r["evidence_value"][:60] + ("..." if len(r["evidence_value"]) > 60 else ""),
      r["recommended_fix"][:60] + ("..." if len(r["recommended_fix"]) > 60 else "")]
     for r in priority],
)

add_para(doc, "Reading:", bold=True)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "P1 compound splitting and P2 digit/word rendering are the two "
    "highest-leverage mechanical fixes. Both are post-processing "
    "candidates that don't require retraining; they touch the highest "
    "fraction of EM-False rows in the manual audit."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "P3 (OOV) and P10 (non-word emissions) are related: they share the "
    "underlying cause of decoder emission without a dictionary "
    "constraint. Constrained decoding over CMU (or g2p-extension of "
    "CMU to cover the OOVs found) would address both at once."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "P4 (open-vowel confusion) is the only category that genuinely "
    "needs acoustic-side work (vowel length is not feature-addressable "
    "by simple rescoring; see Section 9)."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "P5 (casing artefacts) is cosmetic. Already mitigated by reporting "
    "CER separately from WER; mechanical fix is to lower-case all "
    "hypotheses (or upper-case all references) at evaluation time."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "P7 (homophones) and P9 (closed-class grammar detector) are NOT "
    "errors in the model — they reflect the corpus not containing the "
    "particular confusion shape the detector is designed to catch. "
    "Keep the detector armed for production-scale runs; do not invest "
    "in it specifically for this corpus."
)

add_figure(doc, FIGURES / "fig18_priority_ranked.png",
    "Figure 18 — Priority-ranked issues from Step 10. P1 = highest "
    "leverage. Bar length is the rank-order signal; the annotation "
    "shows the per-issue evidence_value.")


# ---------- Update the Reproduction section to mention Stage 4 scripts ----

repro_p = None
paragraphs = list(doc.paragraphs)
for i, p in enumerate(paragraphs):
    if p.style.name == "Heading 1" and p.text.strip() == "Reproduction":
        for q in paragraphs[i + 1:]:
            if q.style.name == "Normal":
                repro_p = q
                break
        break

if repro_p is not None:
    append_text = (" Stage 4 grammar / semantic / consolidated summary: "
                   "step4_grammar_semantic.py (Steps 7 + 8; BERTScore "
                   "uses microsoft/deberta-base-mnli on CPU) and "
                   "stage4_summary.py (Step 10 priority-ranked summary). "
                   "Figures: _make_stage4_figures.py produces "
                   "fig16-fig18. No GPU required.")
    repro_p.add_run(append_text)

doc.save(DOC)
print(f"appended Stage 4 sections (11-13) to {DOC}")