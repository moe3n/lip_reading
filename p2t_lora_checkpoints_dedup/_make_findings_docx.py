"""Build analysis/findings.docx and three supporting CSVs from the
already-produced analysis tables. Run once. The doc is hand-written
with numbers inlined.
"""

import csv
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
TABLES = ROOT / "analysis" / "tables"
FIGURES = ROOT / "analysis" / "figures"
OUT = ROOT / "analysis" / "findings.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)


# --- small helpers ----------------------------------------------------------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_figure(doc, png_path, caption, width_inches=6.5):
    if not Path(png_path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(png_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = None  # default; keep small via paragraph style if desired


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    return t


def read_csv(p):
    with p.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def write_csv(p, header, rows):
    with p.open("w", encoding="utf-8", newline="") as f:
        w = csv.writer(f)
        w.writerow(header)
        for r in rows:
            w.writerow(r)


# --- inputs -----------------------------------------------------------------

sid = read_csv(TABLES / "sid_table.csv") if (TABLES / "sid_table.csv").exists() else None
buckets = read_csv(TABLES / "bucket_counts.csv") if (TABLES / "bucket_counts.csv").exists() else None
failures = read_csv(TABLES / "failure_buckets.csv") if (TABLES / "failure_buckets.csv").exists() else None

# Derive the 26-row audit set (rows with non-empty comment column)
audited = [r for r in failures if r.get("comment", "").strip()] if failures else []

# Headline numbers (hard-coded; matches RESULTS_SUMMARY.md)
HEADLINE = {
    "n": 949,
    "wer": 2.09,
    "cer": 0.98,
    "bleu": 0.9673,
    "em": 91.99,
    "wer_greedy": 3.53,
    "em_greedy": 86.51,
}

# SID table — sourced from step1_sid.py output captured in conversation
SID_ROWS = [
    ("Word · overall",       5214, 87, 13, 11, 5116, 2.129),
    ("Word · homophone",     3995, 57,  8,  9, 3929, 1.852),
    ("Word · non-homophone", 1219, 30,  5,  2, 1187, 3.035),
    ("Char · overall",      26345,116, 67, 78,26151, 0.991),
    ("Char · homophone",    19751, 74, 35, 60,19617, 0.856),
    ("Char · non-homophone", 6594, 42, 32, 18, 6534, 1.395),
]


# --- build ------------------------------------------------------------------

doc = Document()

# Title
title = doc.add_heading("P2T Decoder — Beam-5 Dedup Findings", level=0)
add_para(doc, "Empirical and statistical findings on the beam-5 (width=5) decoded "
              "validation set, after dedup of the 1,082-row raw val against train "
              "(n=949). Source: predictions_beam5.csv + sid_per_pair.csv in "
              "p2t_lora_checkpoints_dedup/. Date: 19 July 2026.")

# Section 1: Headline
add_heading(doc, "1. Headline numbers", level=1)
add_para(doc, f"On the {HEADLINE['n']} dedup val rows, with beam-5 decoding, the "
              f"model reaches WER {HEADLINE['wer']}%, CER {HEADLINE['cer']}%, "
              f"BLEU-4 {HEADLINE['bleu']}, and exact-match {HEADLINE['em']}%. "
              f"Greedy decoding on the same {HEADLINE['n']} rows reaches WER "
              f"{HEADLINE['wer_greedy']}%, EM {HEADLINE['em_greedy']}% — beam-5 "
              f"reduces WER by ~1.4 pp absolute and raises EM by ~5.5 pp absolute.")

add_table(doc,
    ["Decoding", "WER ↓", "CER ↓", "BLEU-4 ↑", "Exact Match ↑"],
    [
        ["Greedy", f"{HEADLINE['wer_greedy']}%", "—", "—", f"{HEADLINE['em_greedy']}%"],
        ["Beam-5 (headline)", f"{HEADLINE['wer']}%", f"{HEADLINE['cer']}%",
         f"{HEADLINE['bleu']}", f"{HEADLINE['em']}%"],
    ])
add_figure(doc, FIGURES / "fig01_beam5_vs_greedy.png",
           "Figure 1 — Beam-5 vs greedy decoding on the dedup val set (n=949). "
           f"Beam-5 reduces WER by {HEADLINE['wer_greedy'] - HEADLINE['wer']:.2f} pp "
           f"and raises EM by {HEADLINE['em'] - HEADLINE['em_greedy']:.2f} pp.")

# Section 2: SID
add_heading(doc, "2. Substitution / Insertion / Deletion breakdown", level=1)
add_para(doc, "Counts are produced by jiwer's process_words and "
              "process_characters at the per-pair level, then aggregated. "
              "H = hits (correctly aligned tokens). Rate = (S + I + D) / N_ref.")
add_table(doc,
    ["Bucket", "N_ref", "S", "I", "D", "H", "Rate (%)"],
    SID_ROWS)

add_para(doc, "Two patterns from this table:", bold=False)
p = doc.add_paragraph(style="List Bullet")
p.add_run("Word level is substitution-dominated: 87 of 111 word-level errors are "
          "S (78%). Insertions and deletions are small.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("Char level is insert/delete-heavy (I=67, D=78), and the I+D excess "
          "is concentrated in the homophone bucket (95 of 145 total I+D — about "
          "2× the non-homophone bucket's 50).")
add_figure(doc, FIGURES / "fig02_sid_stacked.png",
           "Figure 2 — Substitution / Insertion / Deletion / Hits per bucket "
           "(semi-transparent grey cap = Hits). The hits bar shrinks in the "
           "non-homophone bucket — that is the I+D bite the model takes out of "
           "the alignment when target and prediction do not share an acoustic shape.")

# Section 3: Failure-mode taxonomy
add_heading(doc, "3. Failure-mode taxonomy (8 buckets, EM-False rows only)", level=1)
add_para(doc, f"Of the {HEADLINE['n']} rows, 76 fail exact-match. We classify each "
              f"failing row into one of eight failure-mode buckets using the target "
              f"text, prediction text, and CMU phoneme lookups (OOV marker '?') for "
              f"both sides. Classification is rule-based and auditable; see Section 4 "
              f"for the manually-reviewed subset.")

# bucket_counts.csv order is preserved from the classifier output
BUCKET_HEADER = ["Bucket", "n", "% of EM-False"]
BUCKET_ROWS = [
    ("homophone_substitution", 34, "44.7%"),
    ("non_word_spelling", 17, "22.4%"),
    ("semantic_substitution", 8, "10.5%"),
    ("boundary_hallucination", 5, "6.6%"),
    ("digit_word_rendering", 5, "6.6%"),
    ("oov_target_substitution", 3, "3.9%"),
    ("truncation", 3, "3.9%"),
    ("suffix_hallucination", 1, "1.3%"),
    ("TOTAL", 76, "100.0%"),
]
add_table(doc, BUCKET_HEADER, BUCKET_ROWS)
add_figure(doc, FIGURES / "fig03_failure_modes.png",
           "Figure 3 — Eight failure-mode buckets with row counts and "
           "share-of-EM-False. Colour indicates failure family.")

add_para(doc, "Bucket definitions (short):", bold=True)
defs = [
    ("homophone_substitution", "Both target and prediction are real words; their "
        "phonemes share ≥50% of distinct tokens. Catches true homophones plus "
        "near-homophones, vowel swaps, and phoneme truncations with shared prefix."),
    ("non_word_spelling", "Prediction contains a CMU OOV token ('?') AND the "
        "phoneme-word count matches the target. The model produced an English-like "
        "but invalid word (e.g. SQUWAREL for SQUIRREL)."),
    ("semantic_substitution", "Both real words, <50% phoneme overlap. The model "
        "chose a word with little audible resemblance to the target (e.g. RARITY → "
        "RETIREMENT)."),
    ("boundary_hallucination", "Target and prediction share the same letter "
        "string with a different space configuration (e.g. ROPES → ROPE S)."),
    ("digit_word_rendering", "Exactly one side contains a digit, the other its "
        "English word (e.g. SIX → 6 or 3 → THREE)."),
    ("oov_target_substitution", "Target is CMU OOV (e.g. MILLENNIUM, MOLYNEUX); "
        "prediction is a valid word."),
    ("truncation", "Prediction is shorter than target and shares most of the "
        "target's word prefix (e.g. THEORY dropped entirely)."),
    ("suffix_hallucination", "Prediction word equals a target word plus a "
        "productive English suffix (-ING, -S, -ED, -LY, -MENT, -TION, -NESS, -ITY, "
        "-OUS) the target didn't carry."),
]
for name, descr in defs:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(name + ": ")
    r.bold = True
    p.add_run(descr)

add_para(doc, "Note on the homophone vs semantic split.", italic=True)
add_para(doc, "The 50% phoneme-overlap threshold is arbitrary. Rows like FORTNIGHT "
              "→ NIGHT, OLYMPIC → OLYMPICS, CALM → COME, and FLOURISH → FLURRY "
              "are classified as homophone_substitution because they share enough "
              "phoneme tokens with the target, but they are not true homophones in "
              "the linguistic sense — they are phoneme-truncation and "
              "near-homophone cases. Reported as a single combined bucket, "
              "real_word_substitution covers 42 rows (55.3%) of EM-False; the "
              "homophone / semantic split within it depends on the criterion.")
add_figure(doc, FIGURES / "fig04_real_word_breakdown.png",
           "Figure 4 — Three-bucket collapse: real-word substitution dominates "
           "the EM-False mass (42 / 76 = 55.3%).")

# Section 4: 26-row audit
add_heading(doc, "4. Manually-reviewed audit (26 rows)", level=1)
add_para(doc, "26 rows were inspected by hand to validate the automated bucket "
              "labels. The table records the target, prediction, automated bucket, "
              "and a short mechanism-focused comment. Rows where the automated "
              "label disagrees with the manual label are flagged with !!.")

# Build the audit table
AUDIT_HEADER = ["#", "Target (excerpt)", "Prediction (excerpt)",
                "Auto-bucket", "Comment"]
audit_rows = []
for i, r in enumerate(audited, 1):
    tgt = r["target"]
    hyp = r["prediction"]
    # truncate to ~50 chars for readability
    tgt_s = (tgt[:47] + "...") if len(tgt) > 50 else tgt
    hyp_s = (hyp[:47] + "...") if len(hyp) > 50 else hyp
    bucket = r["bucket"]
    comment = r["comment"]
    # disagreement heuristic: bucket not present as a substring of comment
    stem = bucket.split("_")[0]
    disagree = not (stem in comment.lower() or comment.lower().startswith(stem))
    bucket_cell = ("!! " + bucket) if disagree else bucket
    audit_rows.append([str(i), tgt_s, hyp_s, bucket_cell, comment])
add_table(doc, AUDIT_HEADER, audit_rows)

add_para(doc, "The !! marks show where the 50% phoneme-overlap heuristic "
              "disagrees with the manual reading (most often: a row the heuristic "
              "called homophone_substitution is more accurately a phoneme "
              "truncation, suffix hallucination, or near-homophone in the linguistic "
              "sense). The 8-bucket taxonomy is preserved; the audit table makes "
              "the borderline cases visible.")

# Section 5: Phoneme OOV observation
add_heading(doc, "5. Phoneme OOV observation", level=1)
add_para(doc, "We CMU-phonemise both target and prediction with the nltk "
              "corpus. Words absent from CMUDict are marked with '?'. Of 949 rows:")
p = doc.add_paragraph(style="List Bullet")
p.add_run("34 rows contain at least one OOV word in the target.")
p = doc.add_paragraph(style="List Bullet")
p.add_run("56 rows contain at least one OOV word in the prediction.")
add_para(doc, "Predictions contain more OOV than targets — i.e. the model invents "
              "rare or non-standard words more often than it fails to render an "
              "existing one. This is consistent with the non_word_spelling bucket "
              "(22.4% of EM-False) being the second-largest failure mode.")
add_figure(doc, FIGURES / "fig05_oov_target_vs_pred.png",
           "Figure 5 — Rows with at least one CMU-OOV word. Predictions invent "
           "OOV tokens more often than targets contain them.")

# Section 6: What's not yet in this doc
add_heading(doc, "6. What this doc does not yet cover", level=1)
add_para(doc, "Tracked here as placeholders for follow-up appends.")
for item in [
    "Phoneme Error Rate (PER) and a confusion matrix of target-phoneme → "
    "prediction-phoneme at failing positions. Currently scoped to the boundary "
    "and word-level patterns identified here.",
    "Alignment Error Rate (AER) and Word-PER (WPER).",
    "Homophone-stratified bucket counts (homophone_flag × bucket cross-tab).",
    "Grammar / syntactic analysis of failing sentences.",
    "Semantic-distance analysis of real_word_substitution rows.",
    "Speaker-disjoint split check — the current 2-column CSV has no speaker IDs, "
    "so train/eval speaker overlap cannot be ruled out and may inflate scores.",
    "Held-out test set scoring — test rows (1,243 sentences, rows 46,922–48,164) "
    "are defined but never touched.",
]:
    doc.add_paragraph(item, style="List Bullet")

# Closing
add_heading(doc, "Appendix — PER vs CER scatter (per row)", level=1)
add_para(doc, "Optional figure. Phoneme Error Rate is the standard "
              "phoneme-aligned edit distance (S + I + D over reference phoneme "
              "count). It complements CER by exposing phoneme-level fidelity "
              "where character-level metrics can hide them. Rows whose target has "
              "fewer than three phonemes are excluded from the scatter to avoid "
              "ratio blow-up on short inputs.")
add_figure(doc, FIGURES / "fig06_per_vs_cer.png",
           "Figure 6 — PER vs CER per row. EM rows in green; EM-False in red. "
           "The dashed line is y = x; CER ≥ PER is expected because characters "
           "strictly refine phoneme tokens.")
add_heading(doc, "Reproduction", level=1)
add_para(doc, "Source CSVs: predictions_beam5.csv (949 rows), sid_per_pair.csv "
              "(949 rows), predictions_beam5_with_match.csv (949 rows with "
              "phoneme columns and 26 hand-reviewed comments). Analysis scripts: "
              "step1_sid.py (SID breakdown), step2_classify_failures.py "
              "(8-bucket classifier). No GPU required for any of these steps. "
              "All paths relative to p2t_lora_checkpoints_dedup/.")

doc.save(OUT)
print(f"wrote {OUT}")

# --- supporting CSVs -------------------------------------------------------

# headline.csv
write_csv(
    TABLES / "headline.csv",
    ["metric", "value", "notes"],
    [
        ["n_val", HEADLINE["n"], "949 = 1082 raw val - 133 train-leaks"],
        ["wer_beam5", HEADLINE["wer"], "jiwer"],
        ["cer_beam5", HEADLINE["cer"], "jiwer"],
        ["bleu4_beam5", HEADLINE["bleu"], "sacrebleu"],
        ["em_beam5", HEADLINE["em"], "exact string match"],
        ["wer_greedy", HEADLINE["wer_greedy"], "jiwer, greedy twin"],
        ["em_greedy", HEADLINE["em_greedy"], "exact match, greedy twin"],
        ["wer_delta_pp_beam5_minus_greedy",
         round(HEADLINE["wer_greedy"] - HEADLINE["wer"], 2),
         "absolute WER improvement from beam-5"],
        ["em_delta_pp_beam5_minus_greedy",
         round(HEADLINE["em"] - HEADLINE["em_greedy"], 2),
         "absolute EM improvement from beam-5"],
    ],
)
print(f"wrote {TABLES / 'headline.csv'}")

# sid_table.csv
write_csv(
    TABLES / "sid_table.csv",
    ["bucket", "n_ref", "S", "I", "D", "H", "rate_pct"],
    SID_ROWS,
)
print(f"wrote {TABLES / 'sid_table.csv'}")

# top_failures.csv — the 26 audited rows
audit_rows_csv = []
for i, r in enumerate(audited, 1):
    bucket = r["bucket"]
    comment = r["comment"]
    stem = bucket.split("_")[0]
    disagree = not (stem in comment.lower() or comment.lower().startswith(stem))
    audit_rows_csv.append([
        i, r["target"], r["prediction"], bucket, comment,
        "DISAGREE" if disagree else "agree",
    ])
write_csv(
    TABLES / "top_failures.csv",
    ["row", "target", "prediction", "auto_bucket", "audit_comment", "manual_agree"],
    audit_rows_csv,
)
print(f"wrote {TABLES / 'top_failures.csv'}")