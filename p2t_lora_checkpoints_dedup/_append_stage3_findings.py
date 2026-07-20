"""Append Sections 7-10 to findings.docx documenting the Stage 3
extended-metrics results (WPER, AER, per-phoneme drill-down, error-type
breakdown). Inserts before the existing Appendix section.

Stage 3 inputs (already produced by step3_extended_metrics.py):
    analysis/tables/wper_breakdown.csv
    analysis/tables/aer_breakdown.csv
    analysis/tables/per_phoneme_drilldown.csv
    analysis/tables/error_type_breakdown.csv

Stage 3 figures (already produced by _make_stage3_figures.py):
    analysis/figures/fig12_wper_vs_per.png
    analysis/figures/fig13_aer_by_feature.png
    analysis/figures/fig14_per_phoneme_drilldown.png
    analysis/figures/fig15_error_type_breakdown.png

Run once. Idempotent only if you delete the previously-appended sections
first; otherwise duplicates accumulate.
"""

import csv
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
TABLES = ROOT / "analysis" / "tables"
FIGURES = ROOT / "analysis" / "figures"
DOC = ROOT / "analysis" / "findings.docx"


# --- helpers (mirror _make_findings_docx.py) -------------------------------

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_figure(doc, png_path, caption, width_inches=6.5):
    if not Path(png_path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(png_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    return t


def read_csv(p):
    with p.open(encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def fmt_pct(x):
    return f"{float(x):.2f}%"


# --- build -----------------------------------------------------------------

doc = Document(DOC)

# ---------- Section 7: WPER (Step 4) ---------------------------------------

add_heading(doc, "7. Weighted Phoneme Error Rate (WPER)", level=1)
add_para(doc,
    "Step 4 of the extended metrics. WPER is PER with a per-substitution "
    "weight derived from a feature-distance function over ARPAbet phonemes. "
    "We use the heuristic variant from p2t_lora.evaluation.extended_metrics: "
    "substitutions inside the same manner class weight 0.2, same place OR "
    "same voicing weight 0.4, otherwise 1.0. WPER/PER < 1 ⇒ errors are "
    "phonetically close; WPER/PER ≈ 1 ⇒ errors are not feature-adjacent."
)
add_para(doc,
    "Computed on the full 949-row corpus (no per-pair thresholding). S, I, "
    "D are the standard jiwer phoneme counts from word2phone alignment."
)

wper = read_csv(TABLES / "wper_breakdown.csv")
add_table(
    doc,
    ["Group", "n", "PER", "WPER (heuristic)", "WPER / PER"],
    [[r["group"], r["n"], fmt_pct(r["per"]), fmt_pct(r["wper_heuristic"]),
      f"{float(r['wper_per_ratio']):.3f}"] for r in wper],
)

add_para(doc,
    "Reading the table:", bold=True
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "The WPER/PER ratio is 0.930 across every group (overall, homophone, "
    "non-homophone, em_false). It does not differ between homophone and "
    "non-homophone rows, and it does not differ between the high-error "
    "em_false slice and the corpus as a whole."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Interpretation: the substitutions in this model are NOT predominantly "
    "phonetically close. A ratio near 1 means the average substitution costs "
    "as much in WPER as it does in raw PER — feature-distance weighting does "
    "not pull the error down. Mira Fleite's WPER << PER signal (errors "
    "concentrated in manner, recoverable by a feature-aware rescoring pass) "
    "does not appear here. There is no obvious low-hanging fruit from a "
    "feature-aware rerank."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Non-homophone rows have higher absolute PER (2.52% vs 1.48%) but the "
    "same ratio (0.935). The error mix is louder, not closer."
)

add_figure(doc, FIGURES / "fig12_wper_vs_per.png",
    "Figure 12 — PER vs WPER per group. Bars are within ~7% of each other "
    "in every group; the feature-distance weighting barely moves the total.")

# ---------- Section 8: Allophonic Error Rate (AER, Step 5) -----------------

add_heading(doc, "8. Allophonic Error Rate (AER) by feature dimension", level=1)
add_para(doc,
    "Step 5. For each substitution, we look up the place / manner / voicing "
    "of the target and predicted ARPAbet tokens from a hand-coded feature "
    "table and count whether the two share that feature. The rate is "
    "share_count / n_substitutions on the EM-False slice (where "
    "substitutions exist)."
)

aer = read_csv(TABLES / "aer_breakdown.csv")
add_table(
    doc,
    ["Group", "n", "n_subs", "Place %", "Manner %", "Voicing %"],
    [[r["group"], r["n"], r["n_substitutions"],
      fmt_pct(r["place_pct"]), fmt_pct(r["manner_pct"]),
      fmt_pct(r["voicing_pct"])] for r in aer],
)

add_para(doc, "Reading the table:", bold=True)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Manner is the closest feature (47.4% on EM-False overall): about half "
    "of all substitutions preserve the manner class (stop, fricative, "
    "nasal, vowel, etc.). Place is similar (44.7%)."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Voicing is the most distinctive feature (only 26.3% match). When the "
    "model gets voicing wrong, it is wrong by a real articulatory step "
    "(voiced ↔ voiceless), not by a within-class shift."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Homophone rows are MORE feature-adjacent than non-homophone rows on "
    "place and manner (52% / 52% vs 27% / 36%). The homophone bucket's "
    "errors are real near-misses; the non-homophone bucket's errors are "
    "louder and more distant."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Concretely, the model's substitution errors lean toward place-of-"
    "articulation confusion and voicing flips, not toward same-class "
    "vowel-only shifts. This is consistent with the WPER observation "
    "above: errors are not feature-close on average, but the dominant "
    "axis is place + voicing rather than manner-only."
)

add_figure(doc, FIGURES / "fig13_aer_by_feature.png",
    "Figure 13 — AER by feature dimension on the EM-False slice. Manner "
    "and place are similar; voicing is the most distinctive axis.")

# ---------- Section 9: Per-phoneme drill-down (Step 6) --------------------

add_heading(doc, "9. Per-phoneme drill-down", level=1)
add_para(doc,
    "Step 6. We count, for each reference phoneme r, how many times r is "
    "substituted at any aligned position in the corpus, and divide by the "
    "total number of times r appears. The within-phoneme error rate "
    "exposes which reference sounds the model struggles to render. Phonemes "
    "with very few occurrences are flagged as low-confidence."
)

pp = read_csv(TABLES / "per_phoneme_drilldown.csv")
# Sort by within-rate descending, then by ref_occurrences descending
pp_sorted = sorted(pp, key=lambda r: (-float(r["within_phoneme_error_rate"]),
                                      -int(r["ref_occurrences"])))

# Show only "no" (high-confidence) rows in the table; mention the singletons
# separately to keep the table readable
pp_strong = [r for r in pp_sorted if r["low_confidence"] == "no"]
pp_weak = [r for r in pp_sorted if r["low_confidence"] == "yes"]

add_table(
    doc,
    ["Ref phon", "Occurrences", "Subs", "Within-rate", "Top hyp", "Top n"],
    [[r["ref_phon"], r["ref_occurrences"], r["n_substitutions"],
      fmt_pct(r["within_phoneme_error_rate"]),
      r["top_hyp_phon"], r["top_hyp_count"]] for r in pp_strong],
)

add_para(doc,
    f"The remaining {len(pp_weak)} reference phonemes have exactly one "
    f"substitution each and are not shown individually — their within-rate "
    f"is dominated by sampling noise. The high-confidence rows above are "
    f"where the per-phoneme signal is real."
)

add_para(doc, "Reading the table:", bold=True)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Vowels dominate the high-rate rows: AA (14.7%, top confusion → AH), "
    "EY (11.1%, top → L), EH (8.2%, top → ER), and AY (7.4%, top → EY). "
    "All four are open or open-mid vowels. The model's vowel space has "
    "consistent trouble at the open-mid / open endpoints."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Consonant rows have low rates (T 2.2%, S 2.4%, R 3.4%, IH 3.1%, AH "
    "1.1%). Consonant fidelity is high — the model is not losing acoustic "
    "ground on stop or fricative distinctions in aggregate."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "The single most-confused pair in the table is AA → AH (2 hits at "
    "the top). These two vowels are both low and back, differing mainly "
    "in tenseness / length in standard American English. A feature-aware "
    "post-correction step would not help here — AA and AH share place "
    "and manner; only vowel length distinguishes them, and that is "
    "exactly the dimension the model loses."
)

add_figure(doc, FIGURES / "fig14_per_phoneme_drilldown.png",
    "Figure 14 — Per-phoneme drill-down. Left panel: substitution counts "
    "per reference phoneme. Right panel: within-phoneme error rate per "
    "reference phoneme. Open vowels (AA, EH, EY) dominate the high-rate "
    "side; consonants cluster at the low-rate end.")

# ---------- Section 10: Error-type breakdown (bonus) -----------------------

add_heading(doc, "10. Phoneme-level error-type breakdown", level=1)
add_para(doc,
    "Bonus step. Each phoneme-position-level substitution is classified by "
    "looking at the dominant feature distance and whether the predicted "
    "phoneme matches a CMU-dictionary homophone of the target. The fast "
    "variant uses an exact homophone dict lookup only (no near-homophone "
    "scan over the 125k-word CMU corpus); the slower near-homophone pass "
    "is available but omitted here for runtime reasons. The breakdown is "
    "computed on phoneme-position alignments, so a single failing row can "
    "contribute multiple substitutions and multiple labels."
)

et = read_csv(TABLES / "error_type_breakdown.csv")
add_table(
    doc,
    ["Label", "Count (all 949)", "% of all", "Count (EM-False 76)", "% of EM-False"],
    [[r["label"], r["count_all"], fmt_pct(r["pct_all"]),
      r["count_em_false"], fmt_pct(r["pct_em_false"])] for r in et],
)

add_para(doc, "Reading the table:", bold=True)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "On the EM-False slice, 'Other' is the dominant label at 67.1%. This "
    "is consistent with the word-level finding that most failing rows are "
    "compound (multi-word) errors — splits, casing, OOV names, missing "
    "tokens — not a single clean phoneme substitution."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Vowel substitutions are 18.4% of EM-False phoneme errors. This is the "
    "phoneme-level echo of the per-phoneme drill-down (Section 9) where "
    "open vowels (AA, EH, EY) were the highest-rate rows."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Manner-shifts (7.9%), exact homophones (5.3%), and a single "
    "hallucination (1.3%) account for the rest. There are no clear-cut "
    "voicing flips as a standalone category — voicing errors appear "
    "absorbed into 'Other' or 'Manner' depending on which other feature "
    "shifts alongside them."
)
p = doc.add_paragraph(style="List Bullet")
p.add_run(
    "Take-away for downstream work: the model's residual errors are "
    "characterised by (i) open-vowel within-class confusion at the "
    "phoneme level and (ii) compound structural mismatches at the word "
    "level. Neither is addressable by simple feature-aware rescoring; "
    "the structural mismatches in particular point at training-data "
    "or pre-processing issues rather than decoding-time fixes."
)

add_figure(doc, FIGURES / "fig15_error_type_breakdown.png",
    "Figure 15 — Error-type breakdown. The 'Other' bucket dominates the "
    "EM-False slice; 'Vowel' is the largest named phoneme-level category.")

# ---------- Update the Reproduction section to mention Stage 3 scripts ----

# The original Reproduction paragraph is the last Normal paragraph in the
# document. Find it and append the Stage 3 scripts at the end.
repro_p = None
paragraphs = list(doc.paragraphs)
for i, p in enumerate(paragraphs):
    if p.style.name == "Heading 1" and p.text.strip() == "Reproduction":
        for q in paragraphs[i + 1:]:
            if q.style.name == "Normal":
                repro_p = q
                break
        break

if repro_p is not None:
    append_text = (" Stage 3 extended metrics: step3_extended_metrics.py "
                   "(WPER / AER / per-phoneme / error-type) and "
                   "_make_stage3_figures.py (fig12-fig15). Inputs are the "
                   "phoneme columns in predictions_beam5_with_match.csv; "
                   "no GPU required.")
    repro_p.add_run(append_text)

doc.save(DOC)
print(f"appended Stage 3 sections to {DOC}")