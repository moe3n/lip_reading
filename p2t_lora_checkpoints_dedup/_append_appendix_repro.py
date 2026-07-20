"""Append Appendix + Reproduction sections to findings.docx.

These were missing because intermediate appender scripts saved
without them. Adds:

- Appendix — PER vs CER scatter (per row)
  - 1 figure (fig06) embedded
  - 1 table (per_row_punct_vs_cer.csv summary)
- Reproduction
  - Updated text with Stage 4 references
"""
import csv
import os
from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
DOC = ROOT / "analysis" / "findings.docx"
TABLES = ROOT / "analysis" / "tables"
FIGURES = ROOT / "analysis" / "figures"


def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_figure(doc, png_path, caption, width_inches=6.5):
    if not Path(png_path).exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(png_path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True


def add_table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(header):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    return t


doc = Document(DOC)

# ---------- Appendix -------------------------------------------------------

add_heading(doc, "Appendix — PER vs CER scatter (per row)", level=1)
add_para(
    doc,
    "Why CER understates WER: phoneme-confusion pairs often differ in word "
    "count. Insertion/deletion of a single token may be a 1/N WER bump but "
    "1/(N*L) CER bump, where L is the word length. The scatter plot below "
    "shows per-row PER (Phoneme Error Rate, on phoneme sequence) vs CER "
    "(Character Error Rate, on text) for the 949 dedup val rows. The dot "
    "cloud is bounded above the PER=CER line because every word-level "
    "edit forces multiple character edits; a few CER-only outliers are "
    "punctuation and digit insertions that don't break the phoneme "
    "sequence but inflate CER.",
)
add_figure(
    doc,
    FIGURES / "fig06_per_vs_cer.png",
    "Figure 6 — PER (phoneme) vs CER (character) on 949 dedup val rows. "
    "Most points cluster in the 0–0.05 band on both axes; CER understates "
    "WER on substitution-heavy rows.",
)

per_row_csv = TABLES / "per_row_punct_vs_cer.csv"
if per_row_csv.exists():
    rows = []
    with per_row_csv.open(encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        header = reader.fieldnames
        for r in reader:
            rows.append([r[c] for c in header])
    add_para(doc, "Per-row punct vs CER (top 10 worst CER rows):", bold=True)
    if len(rows) > 10:
        rows = rows[:10]
    add_table(doc, header, rows)

# ---------- Reproduction --------------------------------------------------

add_heading(doc, "Reproduction", level=1)
add_para(
    doc,
    "All numbers in this document are derived from "
    "`predictions_beam5.csv` in this same directory (n=949 dedup val rows). "
    "Stage 1 (SID/failure-modes), Stage 2 (homophone bucket + audit), and "
    "Stage 3 (WPER + AER + per-phoneme + error-type) are produced by "
    "`_make_findings_docx.py`, `_append_stage3_findings.py`, and "
    "`_make_stage3_figures.py`. Stage 4 grammar / semantic / consolidated "
    "summary: `step4_grammar_semantic.py` (Steps 7 + 8; BERTScore uses "
    "`microsoft/deberta-base-mnli` on CPU) and `stage4_summary.py` (Step 10 "
    "priority-ranked summary). Figures: `_make_stage4_figures.py` produces "
    "fig16–fig18. No GPU required. Section reorder to keep sections in "
    "numerical order before Appendix is applied by `_reorder_stage4_sections.py`.",
)

doc.save(DOC)
print(f"appended Appendix + Reproduction to {DOC}; size: {os.path.getsize(DOC)}")
