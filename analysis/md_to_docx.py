"""
Convert a markdown report to .docx.

Handles the subset of markdown these reports actually use: headings, paragraphs,
pipe tables, bold spans, and inline code. Written because pandoc is not
installed on this box.

Usage:
    python analysis/md_to_docx.py <input.md> [output.docx]
"""

import os
import re
import sys

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def add_runs(paragraph, text):
    """Split on **bold** and `code`, adding a run per span."""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            paragraph.add_run(part)


def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.style = "Table Grid"
    table.autofit = True

    for j, cell_text in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.bold = True
        run.font.size = Pt(9.5)

    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, row[j] if j < len(row) else "")
            for run in p.runs:
                run.font.size = Pt(9.5)
            # Right-align numeric columns so figures line up down the column.
            val = row[j] if j < len(row) else ""
            if re.fullmatch(r"[\d,.%]+", val.strip()):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()


def convert(md_path, docx_path):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

    for section in doc.sections:
        section.left_margin = section.right_margin = Inches(1.0)

    render_markdown_into(doc, md_path)
    doc.save(docx_path)
    return docx_path


def render_markdown_into(doc, md_path):
    """Render the markdown at md_path into an existing python-docx Document.
    Used both by convert() and by callers that build their own cover page first.
    Image paths in the markdown resolve relative to md_path."""
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Image: ![caption](path) on its own line. Path is resolved relative
        # to the markdown file so figures/ references work.
        img = IMAGE_RE.match(stripped)
        if img:
            caption, rel = img.group(1), img.group(2)
            img_path = rel if os.path.isabs(rel) else os.path.join(os.path.dirname(md_path), rel)
            if os.path.isfile(img_path):
                doc.add_picture(img_path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap.add_run(caption)
                    run.italic = True
                    run.font.size = Pt(9)
            else:
                doc.add_paragraph(f"[missing image: {rel}]")
            i += 1
            continue

        # Table: collect consecutive pipe rows, skipping the separator.
        if is_table_row(line):
            rows = []
            while i < len(lines) and is_table_row(lines[i]):
                if not is_separator(lines[i]):
                    rows.append(split_row(lines[i]))
                i += 1
            if rows:
                add_table(doc, rows)
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        else:
            # Join wrapped lines into one paragraph until a blank line or a
            # structural token. The markdown is hard-wrapped at ~78 chars and
            # those breaks are not meaningful.
            buf = [stripped]
            i += 1
            while i < len(lines):
                nxt = lines[i].strip()
                if (not nxt or nxt.startswith("#") or is_table_row(lines[i])):
                    break
                buf.append(nxt)
                i += 1
            p = doc.add_paragraph()
            add_runs(p, " ".join(buf))
            continue

        i += 1


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("usage: python analysis/md_to_docx.py <input.md> [output.docx]")
    src = sys.argv[1]
    dst = sys.argv[2] if len(sys.argv) > 2 else os.path.splitext(src)[0] + ".docx"
    print(f"Wrote {convert(src, dst)}")
